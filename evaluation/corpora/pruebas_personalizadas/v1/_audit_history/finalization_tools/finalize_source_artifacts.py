#!/usr/bin/env python3
"""Apply the narrowly authorized source corrections for corpus finalization."""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from pypdf import PdfReader, PdfWriter
from pypdf.generic import DecodedStreamObject, NameObject


ROOT = Path(__file__).resolve().parents[1]
LENGTH_NOTE = (
    "La extensión indicada es una guía de formato. No constituye por sí misma "
    "un criterio de evaluación; el trabajo se evalúa por la evidencia y el "
    "razonamiento solicitados."
)


DOCX_REPLACEMENTS: dict[str, list[tuple[str, str]]] = {
    "activity_05_visitas_a_bibliotecas/01_assignment.docx": [
        (
            "Informe de 500 a 900 palabras con una tabla breve. No hace falta una prueba estadística. Debes distinguir resultados calculados, interpretación y datos ausentes.",
            f"Extensión orientativa: informe de 500 a 900 palabras con una tabla breve. No hace falta una prueba estadística. Debes distinguir resultados calculados, interpretación y datos ausentes. {LENGTH_NOTE}",
        )
    ],
    "activity_06_movilidad_estudiantil/01_assignment.docx": [
        (
            "Memo de 500 a 900 palabras. No busques tasas de transporte, accidentes ni emisiones. Se evalúa la calidad del argumento con este dossier, no que exista una única política correcta.",
            f"Extensión orientativa: memo de 500 a 900 palabras. No busques tasas de transporte, accidentes ni emisiones. Se evalúa la calidad del argumento con este dossier, no que exista una única política correcta. {LENGTH_NOTE}",
        )
    ],
    "activity_07_aislamiento_termico/01_assignment.docx": [
        (
            "Un informe de 450 a 800 palabras con tabla, conclusión y comparación cuantitativa.",
            f"Extensión orientativa: un informe de 450 a 800 palabras con tabla, conclusión y comparación cuantitativa. {LENGTH_NOTE}",
        )
    ],
    "activity_08_triage_de_logs/01_assignment.docx": [
        ("10.1.4.22", "192.0.2.44"),
    ],
    "activity_09_renovacion_y_desplazamiento/01_assignment.docx": [
        (
            "Ensayo de 1.000 a 1.600 palabras. Las citas se identifican como Fuente A, B o C. No investigues leyes, autoridades ni procesos urbanos externos. 'Contextualizar' significa situar una afirmación dentro de fechas, propósitos y poblaciones que aparecen en el dossier.",
            f"Extensión orientativa: ensayo de 1.000 a 1.600 palabras. Las citas se identifican como Fuente A, B o C. No investigues leyes, autoridades ni procesos urbanos externos. 'Contextualizar' significa situar una afirmación dentro de fechas, propósitos y poblaciones que aparecen en el dossier. {LENGTH_NOTE}",
        )
    ],
    "activity_10_experimento_onboarding/01_assignment.docx": [
        (
            "report.pdf o report.md de 600 a 1.000 palabras.",
            f"Extensión orientativa: report.pdf o report.md de 600 a 1.000 palabras. {LENGTH_NOTE}",
        )
    ],
    "activity_11_duplicados_en_pagos/01_assignment.docx": [
        ("Extensión esperada", "Extensión orientativa / no evaluable"),
        (
            "El postmortem debe tener entre 600 y 1.000 palabras y el ADR entre 450 y 800. El parche mantiene el máximo de 80 líneas. La extensión debe emplearse en reconstruir estados y fallos, comparar alternativas y explicar recuperación; no en repetir terminología de sistemas distribuidos. Los tres artefactos se valoran como una unidad coherente.",
            f"Como guía de formato, el postmortem puede tener entre 600 y 1.000 palabras y el ADR entre 450 y 800. El parche mantiene el máximo evaluable de 80 líneas. La extensión orientativa puede emplearse en reconstruir estados y fallos, comparar alternativas y explicar recuperación; no en repetir terminología de sistemas distribuidos. Los tres artefactos se valoran como una unidad coherente. {LENGTH_NOTE}",
        ),
        (
            "postmortem.md de 600 a 1.000 palabras: impacto conocido, línea temporal, causa técnica, factores contribuyentes, detección, acciones y límites de evidencia.",
            "postmortem.md (orientación de formato: 600 a 1.000 palabras, no criterio evaluable): impacto conocido, línea temporal, causa técnica, factores contribuyentes, detección, acciones y límites de evidencia.",
        ),
    ],
    "activity_12_clinica_movil/01_assignment.docx": [
        (
            "brief.pdf o brief.md de 650 a 1.100 palabras.",
            f"Extensión orientativa: brief.pdf o brief.md de 650 a 1.100 palabras. {LENGTH_NOTE}",
        )
    ],
}


def replace_docx_text(path: Path, replacements: list[tuple[str, str]]) -> None:
    doc = Document(path)
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)

    for old, new in replacements:
        matches = [p for p in all_paragraphs if p.text == old]
        if len(matches) != 1:
            raise RuntimeError(f"Expected exactly one DOCX match in {path}: {old!r}; found {len(matches)}")
        paragraph = matches[0]
        if len(paragraph.runs) != 1:
            raise RuntimeError(f"Expected a single formatted run in {path}: {old!r}")
        paragraph.runs[0].text = new

    tmp = path.with_suffix(path.suffix + ".finalizing")
    doc.save(tmp)
    os.replace(tmp, path)


def replace_pdf_content(path: Path) -> None:
    reader = PdfReader(path)
    writer = PdfWriter()
    writer.clone_document_from_reader(reader)
    old = b"concentra en el grupo cercano: 15 de 20, frente a 7 de 25 en distancia media y 4 de 15 en distancia larga. Esa"
    new = b"concentra en el grupo cercano: 15 de 20, frente a 8 de 25 en distancia media y 3 de 15 en distancia larga. Esa"
    replacements = 0
    for page in writer.pages:
        contents = page.get_contents()
        if contents is None:
            continue
        data = contents.get_data()
        count = data.count(old)
        if count:
            data = data.replace(old, new)
            stream = DecodedStreamObject()
            stream.set_data(data)
            page[NameObject("/Contents")] = writer._add_object(stream)
            replacements += count
    if replacements != 1:
        raise RuntimeError(f"Expected one PDF content replacement in {path}; found {replacements}")
    tmp = path.with_suffix(path.suffix + ".finalizing")
    with tmp.open("wb") as handle:
        writer.write(handle)
    os.replace(tmp, path)


def main() -> None:
    for relative, replacements in DOCX_REPLACEMENTS.items():
        replace_docx_text(ROOT / relative, replacements)
    replace_pdf_content(
        ROOT
        / "activity_06_movilidad_estudiantil/submissions/submission_01_strong.pdf"
    )


if __name__ == "__main__":
    main()
