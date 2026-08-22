"""Build the deterministic synthetic DOCX pack used by semantic qualification.

The documents are inert fixtures: no links, macros, fields, embedded objects,
student data, or provider content.  They intentionally exercise the same DOCX
parser boundary as the product.
"""

from __future__ import annotations

import argparse
from datetime import UTC, datetime
from pathlib import Path
import tempfile
import zipfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT = (
    ROOT / "tests/fixtures/openai_evals/v3/document_shaped_cache_case"
)
FIXED_TIME = datetime(2026, 8, 13, 12, 0, tzinfo=UTC)
NAVY = "17324D"
BLUE = "356B87"
PALE_BLUE = "E8F1F5"
PALE_GRAY = "F3F5F7"
WHITE = "FFFFFF"


def _set_cell_shading(cell: object, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()  # type: ignore[attr-defined]
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_repeat_table_header(row: object) -> None:
    properties = row._tr.get_or_add_trPr()  # type: ignore[attr-defined]
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def _configure_document(title: str) -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.05

    for style_name, size, color in (
        ("Title", 20, NAVY),
        ("Heading 1", 13.5, NAVY),
        ("Heading 2", 11.5, BLUE),
    ):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.text = "CVA  |  PAQUETE CANÓNICO SINTÉTICO"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Arial"
        run.font.size = Pt(7.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(BLUE)

    footer = section.footer.paragraphs[0]
    footer.text = "SYNTHETIC_ONLY_NO_STUDENT_DATA  •  revisión 2026-08-13"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.name = "Arial"
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor.from_string(BLUE)

    properties = document.core_properties
    properties.title = title
    properties.subject = "HARNESS_SEMANTIC_REMEDIATION"
    properties.author = "CVA Harness Review"
    properties.last_modified_by = "CVA Harness Review"
    properties.keywords = "synthetic, qualification, cache invalidation"
    properties.comments = "Inert synthetic fixture; no student data."
    properties.created = FIXED_TIME
    properties.modified = FIXED_TIME
    return document


def _add_title(document: Document, title: str, subtitle: str) -> None:
    paragraph = document.add_paragraph(style="Title")
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run(title)
    sub = document.add_paragraph()
    sub.paragraph_format.space_after = Pt(10)
    run = sub.add_run(subtitle)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(BLUE)


def _add_metadata(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(1.55)
    table.columns[1].width = Inches(5.75)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(1.55)
        cells[1].width = Inches(5.75)
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_shading(cells[0], PALE_BLUE)
        cells[0].paragraphs[0].add_run(label).bold = True
        cells[1].paragraphs[0].add_run(value)
        for cell in cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8.5)


def _add_callout(document: Document, label: str, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    _set_cell_shading(cell, PALE_BLUE)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    lead = paragraph.add_run(f"{label}: ")
    lead.bold = True
    paragraph.add_run(text)


def _add_grid(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    _set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        _set_cell_shading(cell, NAVY)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            if row_index % 2:
                _set_cell_shading(cells[index], PALE_GRAY)
            cells[index].paragraphs[0].add_run(value)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8)


def _assignment_document() -> Document:
    document = _configure_document(
        "Actividad sintética: trazabilidad de invalidación de caché"
    )
    _add_title(
        document,
        "Trazabilidad de invalidación de caché",
        "Consigna oficial • caso CVA-CACHE-01",
    )
    _add_metadata(
        document,
        [
            ("Clasificación", "SYNTHETIC_ONLY_NO_STUDENT_DATA"),
            ("Tiempo", "5 minutos"),
            ("Formato", "Respuesta abierta breve"),
            ("Materiales", "Sólo el entregable sintético autorizado"),
        ],
    )
    document.add_heading("Propósito", level=1)
    document.add_paragraph(
        "Resultado de aprendizaje: explicar cómo un cambio en la fuente "
        "invalida una entrada almacenada y obliga a recalcular un resultado "
        "verificable."
    )
    document.add_heading("Consigna", level=1)
    document.add_paragraph(
        "Entrega una explicación breve que conecte, en este orden, el cambio "
        "de la fuente, la invalidación de la entrada de caché, la nueva "
        "consulta y el recálculo. Justifica por qué reutilizar la entrada "
        "anterior podría devolver un resultado obsoleto. Usa únicamente la "
        "evidencia incluida en tu entregable."
    )
    document.add_heading("Producto esperado", level=1)
    document.add_paragraph(
        "Un párrafo explicativo y una traza mínima que vinculen la regla de "
        "invalidación con el resultado observado después del cambio de fuente."
    )
    _add_callout(
        document,
        "Límite inferencial",
        "No se pide inferir lenguaje de implementación, rendimiento, "
        "concurrencia ni la causa externa del cambio de fuente.",
    )
    return document


def _rubric_document() -> Document:
    document = _configure_document(
        "Rúbrica sintética: explicación causal de invalidación de caché"
    )
    _add_title(
        document,
        "Rúbrica de evaluación",
        "Documento oficial • caso CVA-CACHE-01",
    )
    _add_metadata(
        document,
        [
            ("Escala", "0–3"),
            ("Criterio", "criterion_1"),
            ("Peso", "100 % (grading_weight = 1.0)"),
            ("Ajuste", "Alta verificabilidad desde evidencia localizada"),
        ],
    )
    document.add_heading("Criterio único", level=1)
    document.add_paragraph(
        "Explicación causal de invalidación de caché: relaciona el cambio de "
        "fuente, la invalidación de la entrada, la nueva consulta y el "
        "recálculo, y explica por qué esa secuencia evita devolver un resultado "
        "obsoleto."
    )
    document.add_heading("Observables", level=2)
    for text in (
        "Ordena la secuencia causal completa.",
        "Vincula la invalidación con la prevención de resultados obsoletos.",
        "Distingue lo sustentado por la traza de lo que no puede inferirse.",
    ):
        document.add_paragraph(text, style="List Bullet")
    document.add_heading("Niveles", level=1)
    _add_grid(
        document,
        ["Nivel", "Descriptor"],
        [
            ["3 • Completo", "Explica la secuencia y el riesgo de obsolescencia con evidencia exacta y sin ampliar la fuente."],
            ["2 • Suficiente", "Explica el núcleo del mecanismo y vincula invalidación con recálculo."],
            ["1 • Parcial", "Menciona pasos correctos, pero omite el vínculo causal o el riesgo evitado."],
            ["0 • No evidenciado", "No presenta una explicación sustentada por el entregable."],
        ],
    )
    return document


def _sufficient_submission_document() -> Document:
    document = _configure_document(
        "Entregable sintético suficiente: invalidación de caché"
    )
    _add_title(
        document,
        "Informe técnico breve",
        "Entregable sintético suficiente • submission_cache_sufficient",
    )
    _add_callout(
        document,
        "Declaración",
        "Contenido enteramente sintético; no representa a una persona ni "
        "contiene datos estudiantiles.",
    )
    document.add_heading("Mecanismo implementado", level=1)
    document.add_paragraph(
        "Cuando cambia el hash de la fuente, el flujo invalida la entrada de "
        "caché antes de la siguiente consulta. Si conservara la entrada "
        "anterior, la consulta podría devolver un resultado calculado con la "
        "versión previa y, por tanto, obsoleto. Al no encontrar una entrada "
        "válida, la nueva consulta recalcula el valor desde la fuente "
        "actualizada."
    )
    document.add_heading("Prueba sintética", level=1)
    document.add_paragraph(
        "La prueba registra 42 para la fuente v1, cambia la fuente a v2, "
        "ejecuta una nueva consulta y obtiene 57, el valor calculado desde v2."
    )
    document.add_heading("Alcance", level=1)
    document.add_paragraph(
        "La traza demuestra el comportamiento localizado. No permite inferir "
        "rendimiento, consistencia concurrente ni lenguaje de implementación."
    )
    return document


def _insufficient_submission_document() -> Document:
    document = _configure_document(
        "Entregable sintético insuficiente: invalidación de caché"
    )
    _add_title(
        document,
        "Informe técnico incompleto",
        "Entregable sintético insuficiente • submission_cache_insufficient",
    )
    _add_callout(
        document,
        "Declaración",
        "Contenido enteramente sintético; no representa a una persona ni "
        "contiene datos estudiantiles.",
    )
    document.add_heading("Mecanismo descrito", level=1)
    document.add_paragraph(
        "Cuando cambia el hash de la fuente, el flujo invalida la entrada de "
        "caché antes de la siguiente consulta."
    )
    document.add_heading("Prueba sintética", level=1)
    document.add_paragraph(
        "La prueba cambia la fuente de v1 a v2. No registra el valor devuelto "
        "por la segunda consulta ni identifica qué riesgo evita la invalidación."
    )
    return document


def _write_deterministic_docx(document: Document, destination: Path) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temporary:
        temporary_path = Path(temporary.name)
    try:
        document.save(temporary_path)
        with zipfile.ZipFile(temporary_path, "r") as source:
            entries = [(info.filename, source.read(info)) for info in source.infolist()]
        with zipfile.ZipFile(
            destination,
            "w",
            compression=zipfile.ZIP_DEFLATED,
            compresslevel=9,
        ) as target:
            # Preserve the canonical OPC part order emitted by python-docx;
            # common libmagic databases use that order to identify OOXML.
            for name, payload in entries:
                info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
                # Keep the OPC discriminator visible to common libmagic
                # databases while compressing every other part.
                info.compress_type = (
                    zipfile.ZIP_STORED
                    if name == "[Content_Types].xml"
                    else zipfile.ZIP_DEFLATED
                )
                target.writestr(info, payload)
    finally:
        temporary_path.unlink(missing_ok=True)


def build_documents(output_dir: Path) -> list[Path]:
    documents = {
        "official_assignment.docx": _assignment_document(),
        "official_rubric.docx": _rubric_document(),
        "submission_sufficient.docx": _sufficient_submission_document(),
        "submission_insufficient.docx": _insufficient_submission_document(),
    }
    paths: list[Path] = []
    for filename, document in documents.items():
        destination = output_dir / filename
        _write_deterministic_docx(document, destination)
        paths.append(destination)
    return paths


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    for path in build_documents(args.output_dir.resolve()):
        print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
